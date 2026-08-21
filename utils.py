from pypdf import PdfReader
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import re
from typing import List, Dict

def extract_text_from_pdf(file) -> str:
    reader = PdfReader(file)
    text = ""
    for page in reader.pages:
        text += (page.extract_text() or "") + "\n"
    return text.strip()

def extract_text_from_docx(file) -> str:
    doc = Document(file)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

def extract_resume_text(uploaded_file) -> str:
    if uploaded_file is None:
        return ""
    name = uploaded_file.name.lower()
    if name.endswith(".pdf"):
        return extract_text_from_pdf(uploaded_file)
    elif name.endswith(".docx"):
        return extract_text_from_docx(uploaded_file)
    return uploaded_file.read().decode("utf-8", errors="ignore")

def simple_keyword_extract(text: str, top_n: int = 40) -> List[str]:
    """Very simple deterministic keyword extractor (no external NLP)."""
    text = text.lower()
    # Remove common stop words
    stop = set("""
        a an the and or but if in on at to for of with by from as is are was were be been being
        this that these those it its he she they we you i me my our your their what which who
        how when where why will would can could should may might must shall about into through
        during before after above below between under again further then once here there all
        any both each few more most other some such no nor not only own same so than too very
        s t can will just don should now
    """.split())
    words = re.findall(r'\b[a-z][a-z0-9+#.\-]{2,}\b', text)
    freq = {}
    for w in words:
        if w not in stop and not w.isdigit():
            freq[w] = freq.get(w, 0) + 1
    sorted_words = sorted(freq.items(), key=lambda x: x[1], reverse=True)
    return [w for w, _ in sorted_words[:top_n]]

def create_ats_docx(name: str, summary: str, skills: List[str], experience: str, education: str) -> bytes:
    """Create a clean single-column ATS-friendly DOCX."""
    doc = Document()
    
    # Name
    title = doc.add_paragraph()
    run = title.add_run(name or "Your Name")
    run.bold = True
    run.font.size = Pt(18)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Summary
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(summary)
    
    # Skills
    doc.add_heading("Skills", level=1)
    doc.add_paragraph(", ".join(skills))
    
    # Experience
    doc.add_heading("Experience", level=1)
    for line in experience.split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())
    
    # Education
    doc.add_heading("Education", level=1)
    doc.add_paragraph(education)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()